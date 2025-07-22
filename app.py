from flask import Flask, render_template, request, send_file, redirect, url_for, flash, jsonify
import os
import tempfile
import whisper
from docx import Document
import math
from pyannote.audio import Pipeline
from dotenv import load_dotenv
import psutil
import gc

app = Flask(__name__)
app.secret_key = 'supersecretkey'
UPLOAD_FOLDER = tempfile.gettempdir()
ALLOWED_EXTENSIONS = {'mp3', 'wav', 'm4a', 'ogg', 'flac'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # 500 MB máximo

# Cargar modelo Whisper una sola vez
print("Cargando modelo Whisper...")
# Modelos disponibles: tiny (39M), base (74M), small (244M), medium (769M), large-v2 (1550M)
# Para Intel Iris Xe, usamos medium que ofrece buen balance calidad/rendimiento
import torch
device = "cuda" if torch.cuda.is_available() else "cpu"
print(f"Usando dispositivo: {device}")
print("Nota: Intel Iris Xe no soporta CUDA, usando CPU optimizado")

# Usar modelo base para mejor eficiencia de memoria
# Para archivos muy largos, es mejor usar un modelo más pequeño
model = whisper.load_model('base', device=device)
print("Modelo Whisper 'base' cargado - optimizado para eficiencia de memoria.")

# Cargar pipeline de diarización de pyannote (requiere token HuggingFace)
load_dotenv()
HUGGINGFACE_TOKEN = os.environ.get('HUGGINGFACE_TOKEN')
pipeline = None
if HUGGINGFACE_TOKEN:
    try:
        print("Cargando pipeline de pyannote...")
        pipeline = Pipeline.from_pretrained('pyannote/speaker-diarization-3.1', use_auth_token=HUGGINGFACE_TOKEN)
        print("Pipeline de pyannote cargado exitosamente.")
    except Exception as e:
        print(f"Error al cargar pipeline de pyannote: {type(e).__name__}: {str(e)}")
        pipeline = None
else:
    print("No se encontró HUGGINGFACE_TOKEN en las variables de entorno.")

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/', methods=['GET'])
def index():
    return redirect(url_for('dashboard'))

@app.route('/dashboard')
def dashboard():
    return render_template('dashboard.html')

@app.route('/upload', methods=['POST'])
def upload():
    if 'audiofile' not in request.files:
        return jsonify({'success': False, 'message': 'No se seleccionó ningún archivo'}), 400
    file = request.files['audiofile']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'No se seleccionó ningún archivo'}), 400
    if file and allowed_file(file.filename):
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(filepath)
        if os.path.getsize(filepath) == 0:
            return jsonify({'success': False, 'message': 'El archivo está vacío. Por favor, selecciona un archivo de audio válido.'}), 400
        
        # Obtener información del archivo
        file_size_mb = os.path.getsize(filepath) / (1024 * 1024)
        file_info = f"Tamaño: {file_size_mb:.2f} MB"
        
        # Intentar obtener duración del audio
        try:
            import subprocess
            result = subprocess.run(['ffprobe', '-v', 'error', '-show_entries', 'format=duration', '-of', 'default=noprint_wrappers=1:nokey=1', filepath], 
                                  capture_output=True, text=True)
            if result.stdout.strip():
                duration_seconds = float(result.stdout.strip())
                duration_minutes = duration_seconds / 60
                file_info += f" - Duración: {duration_minutes:.1f} minutos"
        except:
            pass
        # Validación extra: intentar abrir el archivo con ffmpeg para asegurarse que es audio válido
        import subprocess
        try:
            result = subprocess.run([
                'ffmpeg', '-v', 'error', '-i', filepath, '-f', 'null', '-'
            ], capture_output=True, text=True)
            if result.stderr:
                return jsonify({'success': False, 'message': 'El archivo no es un audio válido o está dañado.'}), 400
        except FileNotFoundError:
            # FFmpeg no está instalado, saltar validación
            print("Advertencia: FFmpeg no está instalado. Saltando validación de audio.")
        except Exception as e:
            print(f"Error en validación de audio: {e}")
            # Continuar sin fallar si hay otros errores
        return jsonify({'success': True, 'filename': file.filename, 'message': f'Archivo subido exitosamente. {file_info}'})
    else:
        return jsonify({'success': False, 'message': 'Formato de archivo no permitido'}), 400

@app.route('/transcribe', methods=['POST'])
def transcribe():
    data = request.get_json()
    filename = data.get('filename')
    idioma = data.get('idioma', 'es')
    num_speakers = int(data.get('num_speakers', 2))
    franjas = data.get('franjas', False)
    if not filename:
        return jsonify({'success': False, 'message': 'No se recibió el nombre del archivo.'}), 400
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(filepath):
        return jsonify({'success': False, 'message': 'El archivo no existe en el servidor.'}), 400
    
    # Verificar tamaño del archivo
    file_size = os.path.getsize(filepath) / (1024 * 1024)  # En MB
    print(f"Tamaño del archivo: {file_size:.2f} MB")
    
    # Obtener duración del audio
    duration_seconds = 0
    try:
        import subprocess
        result = subprocess.run(['ffprobe', '-v', 'error', '-show_entries', 'format=duration', '-of', 'default=noprint_wrappers=1:nokey=1', filepath], 
                              capture_output=True, text=True)
        duration_seconds = float(result.stdout.strip())
        duration_minutes = duration_seconds / 60
        print(f"Duración del audio: {duration_minutes:.2f} minutos")
        
        # Sin límites de tiempo - procesará archivos de cualquier duración
        print(f"Procesando audio de {duration_minutes:.1f} minutos. Esto puede tardar mucho tiempo, por favor espera...")
    except:
        print("No se pudo obtener la duración del audio")
    
    try:
        # Monitorear memoria disponible
        memory = psutil.virtual_memory()
        print(f"Memoria disponible: {memory.available / (1024**3):.2f} GB de {memory.total / (1024**3):.2f} GB")
        
        # Si hay menos de 2GB disponibles, forzar recolección de basura
        if memory.available < 2 * 1024**3:
            gc.collect()
            memory = psutil.virtual_memory()
            print(f"Después de limpieza: {memory.available / (1024**3):.2f} GB disponibles")
        
        # Para archivos grandes (>30 minutos), usar procesamiento por chunks
        CHUNK_DURATION = 300  # 5 minutos por chunk
        use_chunks = duration_seconds > 1800  # Más de 30 minutos
        
        if use_chunks and duration_seconds > 0:
            print(f"Audio largo detectado. Procesando en chunks de {CHUNK_DURATION/60} minutos...")
            
            # Para archivos muy largos, hacer diarización opcional o por chunks
            if duration_seconds > 3600:  # Más de 1 hora
                print("ADVERTENCIA: Audio muy largo (>1 hora). Procesando sin diarización para evitar problemas de memoria.")
                print("Nota: La identificación de hablantes no estará disponible para optimizar el uso de memoria.")
                diarization = None
            else:
                # Intentar diarización con manejo de errores
                if pipeline is None:
                    print("Diarización deshabilitada: No se encontró token de HuggingFace")
                    diarization = None
                else:
                    try:
                        print(f"Iniciando diarización con {num_speakers} hablantes...")
                        # Liberar memoria antes de diarización
                        gc.collect()
                        diarization = pipeline({'audio': filepath}, num_speakers=num_speakers)
                        print("Diarización completada.")
                    except Exception as e:
                        print(f"Error en diarización: {e}")
                        print("Continuando sin diarización...")
                        diarization = None
            
            # Procesar audio en chunks para la transcripción
            import subprocess
            import uuid
            all_segments = []
            num_chunks = int(duration_seconds / CHUNK_DURATION) + 1
            
            for i in range(num_chunks):
                start_time = i * CHUNK_DURATION
                chunk_file = os.path.join(app.config['UPLOAD_FOLDER'], f'chunk_{uuid.uuid4()}.wav')
                
                try:
                    # Extraer chunk usando ffmpeg
                    print(f"Procesando chunk {i+1}/{num_chunks} (minuto {start_time/60:.1f} a {(start_time+CHUNK_DURATION)/60:.1f})...")
                    subprocess.run([
                        'ffmpeg', '-i', filepath, '-ss', str(start_time), 
                        '-t', str(CHUNK_DURATION), '-acodec', 'pcm_s16le', 
                        '-ac', '1', '-ar', '16000', chunk_file, '-y'
                    ], check=True, capture_output=True)
                    
                    # Transcribir chunk
                    result = model.transcribe(
                        chunk_file,
                        fp16=False,
                        language=idioma,
                        word_timestamps=True,
                        task='transcribe',
                        temperature=0,
                        compression_ratio_threshold=2.4,
                        logprob_threshold=-1.0,
                        no_speech_threshold=0.6,
                        condition_on_previous_text=False,  # No usar contexto previo entre chunks
                        initial_prompt="Transcripción precisa y detallada del audio en español.",
                        verbose=False
                    )
                    
                    # Ajustar timestamps relativos al archivo original
                    for seg in result.get('segments', []):
                        seg['start'] += start_time
                        seg['end'] += start_time
                        all_segments.append(seg)
                    
                    # Limpiar chunk temporal
                    os.remove(chunk_file)
                    
                except Exception as e:
                    print(f"Error procesando chunk {i+1}: {e}")
                    if os.path.exists(chunk_file):
                        os.remove(chunk_file)
            
            segments = all_segments
            print(f"Transcripción completada. {len(segments)} segmentos procesados.")
            
        else:
            # Procesamiento normal para archivos más pequeños
            # Diarización con pyannote
            if pipeline is None:
                return jsonify({'success': False, 'message': 'pyannote.audio requiere un HuggingFace Token. Por favor, configura la variable de entorno HUGGINGFACE_TOKEN.'}), 500
            print(f"Iniciando diarización con {num_speakers} hablantes...")
            diarization = pipeline({'audio': filepath}, num_speakers=num_speakers)
            print("Diarización completada.")
            # Transcripción con Whisper - Configuración optimizada para máxima calidad
            print(f"Iniciando transcripción con Whisper base...")
            result = model.transcribe(
                filepath, 
                fp16=False,  # Usar precisión completa
                language=idioma,
                word_timestamps=True,
                task='transcribe',  # No traducir, solo transcribir
                temperature=0,  # Más determinista, menos alucinaciones
                compression_ratio_threshold=2.4,  # Filtrar segmentos de baja calidad
                logprob_threshold=-1.0,  # Filtrar tokens de baja probabilidad
                no_speech_threshold=0.6,  # Mejor detección de silencio
                condition_on_previous_text=True,  # Usar contexto previo
                initial_prompt="Transcripción precisa y detallada del audio en español.",  # Guiar al modelo
                verbose=False  # Sin progreso en terminal para evitar spam
            )
            segments = result.get('segments', [])
            print(f"Transcripción completada. {len(segments)} segmentos procesados.")
        # Asignar segmentos a hablantes
        speaker_segments = []
        for seg in segments:
            if diarization is not None:
                # Buscar el hablante que más se solapa con el segmento
                best_speaker = None
                best_overlap = 0
                for turn, _, speaker in diarization.itertracks(yield_label=True):
                    # turn es un objeto Segment con atributos .start y .end
                    overlap = max(0, min(seg['end'], turn.end) - max(seg['start'], turn.start))
                    if overlap > best_overlap:
                        best_overlap = overlap
                        best_speaker = speaker
                if best_speaker is None:
                    best_speaker = "Hablante"
            else:
                # Sin diarización, todos los segmentos son del mismo hablante
                best_speaker = "Hablante"
            speaker_segments.append({'speaker': best_speaker, 'text': seg['text'], 'start': seg['start'], 'end': seg['end']})
        # Agrupar por franjas si el usuario lo desea
        text_content = ''
        if franjas:
            grouped = []
            current = []
            current_start = 0
            for seg in speaker_segments:
                if seg['start'] - current_start >= 600 and current:
                    grouped.append(current)
                    current = []
                    current_start = seg['start']
                current.append(seg)
            if current:
                grouped.append(current)
            for i, group in enumerate(grouped):
                start_min = math.floor(group[0]['start'] / 60)
                end_min = math.ceil(group[-1]['end'] / 60)
                text_content += f"\n--- Franja {i+1}: minuto {start_min} a {end_min} ---\n"
                for seg in group:
                    text_content += f"{seg['speaker']}: {seg['text']}\n"
        else:
            for seg in speaker_segments:
                text_content += f"{seg['speaker']}: {seg['text']}\n"
        # Guardar TXT
        txt_path = filepath + '.txt'
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        # Guardar DOCX
        doc = Document()
        doc.add_heading('Transcripción de audio', 0)
        if franjas:
            for i, group in enumerate(grouped):
                start_min = math.floor(group[0]['start'] / 60)
                end_min = math.ceil(group[-1]['end'] / 60)
                doc.add_heading(f'Franja {i+1}: minuto {start_min} a {end_min}', level=1)
                for seg in group:
                    doc.add_paragraph(f"{seg['speaker']}: {seg['text']}")
        else:
            for seg in speaker_segments:
                doc.add_paragraph(f"{seg['speaker']}: {seg['text']}")
        docx_path = filepath + '.docx'
        doc.save(docx_path)
        return jsonify({
            'success': True,
            'text': text_content,
            'download_url': f'/download?filename={filename}.txt',
            'download_docx': f'/download?filename={filename}.docx'
        })
    except Exception as e:
        error_type = type(e).__name__
        error_msg = str(e)
        print(f"Error detallado: {error_type}: {error_msg}")
        import traceback
        traceback.print_exc()
        
        # Mensajes de error más específicos para el usuario
        if "CUDA" in error_msg or "GPU" in error_msg:
            user_msg = "Error: No se detectó GPU compatible. La diarización requiere mucha memoria. Intenta con un audio más corto."
        elif "memory" in error_msg.lower() or "killed" in error_msg.lower():
            user_msg = "Error: Memoria insuficiente. El archivo es demasiado grande. Para archivos de más de 2 horas, considera dividirlo en partes más pequeñas."
        elif "pyannote" in error_msg:
            user_msg = f"Error con pyannote: {error_msg}"
        elif error_type == "RuntimeError" and "out of memory" in error_msg.lower():
            user_msg = "Error: Se agotó la memoria del sistema. Intenta con un archivo más pequeño o cierra otras aplicaciones."
        else:
            user_msg = f"Error durante la transcripción: {error_type} - {error_msg}"
            
        return jsonify({'success': False, 'message': user_msg}), 500

@app.route('/download')
def download():
    filename = request.args.get('filename')
    if not filename:
        return 'Archivo no especificado', 400
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(filepath):
        return 'Archivo no encontrado', 404
    ext = os.path.splitext(filename)[1]
    download_name = 'transcripcion' + ext
    return send_file(filepath, as_attachment=True, download_name=download_name)

@app.route('/favicon.ico')
def favicon():
    return '', 204

if __name__ == '__main__':
    # Configuración para manejar archivos grandes - sin debug para evitar reinicios
    print("=== TranscribeAI Pro iniciado ===")
    print("Accede a la aplicación en: http://127.0.0.1:5000")
    print("Presiona Ctrl+C para detener")
    app.run(debug=False, threaded=True, host='127.0.0.1', port=5000)
